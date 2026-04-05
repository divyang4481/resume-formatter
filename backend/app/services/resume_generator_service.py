import io
import json
import re
import logging
from typing import Any, Dict, List, Optional, Tuple

from docxtpl import DocxTemplate, RichText
from docx import Document

from app.services.template_render_planner import TemplateRenderPlanner
from app.services.section_render_engine import SectionRenderEngine

logger = logging.getLogger(__name__)


class ResumeGeneratorService:
    """
    Handles the physical generation of resume documents (DOCX rendering).

    Responsibilities:
      - Normalize template markers (<<…>>, {{…}}, [[…]]) via prepare_document_markers
      - Build render_context via TemplateRenderPlanner + SectionRenderEngine
      - Call docxtpl.render()
      - NOT responsible for AI data generation or harmonization
    """

    def __init__(self, llm=None) -> None:
        self._planner = TemplateRenderPlanner()
        self._render_engine = SectionRenderEngine(llm=llm)


    def render_formatted_document(
        self,
        template_bytes: bytes,
        resume_data: Dict[str, Any],
        expected_fields: Optional[str] = "",
        field_extraction_manifest: Optional[List[Dict[str, Any]]] = None,
        job_id: Optional[str] = None,
    ) -> Tuple[bytes, str]:
        """
        Renders the final DOCX by:
          1. Normalizing template markers
          2. Building a per-field render plan (planner)
          3. Rendering each field via the section engine
          4. Calling docxtpl render
        Returns: (docx_bytes, linearized_data_json)
        """
        try:
            template_stream = io.BytesIO(template_bytes)

            # ── 1. Parse inputs ────────────────────────────────────────────
            expected_fields_list = [
                f.strip() for f in (expected_fields or "").split(",") if f.strip()
            ]

            manifest_list: List[Dict[str, Any]] = []
            if field_extraction_manifest:
                if isinstance(field_extraction_manifest, str):
                    try:
                        manifest_list = json.loads(field_extraction_manifest)
                    except Exception:
                        manifest_list = []
                elif isinstance(field_extraction_manifest, list):
                    manifest_list = [
                        (m.model_dump() if hasattr(m, "model_dump") else
                         m.dict()       if hasattr(m, "dict")       else m)
                        for m in field_extraction_manifest
                        if m is not None
                    ]

            # ── 2. Normalize template markers ──────────────────────────────
            processed_template_stream = self.prepare_document_markers(
                template_stream, expected_fields_list, manifest_list
            )

            # ── 3. Build render plan ───────────────────────────────────────
            render_plan_list = self._planner.build_plan(
                expected_fields=expected_fields_list,
                field_manifest=manifest_list,
                resume_data=resume_data,
            )
            plan_map: Dict[str, Dict] = {p["fieldname"]: p for p in render_plan_list}

            # ── 4. Render each field via section engine ────────────────────
            processed_resume_data: Dict[str, Any] = {}
            linearized_data: Dict[str, str] = {}

            # Build manifest lookup for AI guidance per field
            manifest_by_fieldname: Dict[str, Dict] = {
                m.get("fieldname", ""): m
                for m in manifest_list
                if isinstance(m, dict) and m.get("fieldname")
            }

            for key, value in resume_data.items():
                safe_key = key.lower().strip().replace(" ", "_").replace(":", "")
                plan = plan_map.get(safe_key) or plan_map.get(key) or {
                    "fieldname": safe_key,
                    "render_mode": self._planner._infer_render_mode(value),
                    "schema_hint": {},
                }
                manifest_item = manifest_by_fieldname.get(safe_key) or manifest_by_fieldname.get(key)
                try:
                    docx_val, cvml_val = self._render_engine.render_field(
                        field_name=safe_key,
                        value=value,
                        render_plan=plan,
                        manifest_item=manifest_item,
                        job_id=job_id,
                    )
                    processed_resume_data[safe_key] = docx_val
                    linearized_data[safe_key] = cvml_val
                except Exception as field_err:
                    logger.warning(f"Could not render field '{safe_key}': {field_err}. Using raw value.")
                    processed_resume_data[safe_key] = value
                    linearized_data[safe_key] = str(value)

            # Also populate original-case keys for backward compatibility in render context
            for key, value in resume_data.items():
                if key not in processed_resume_data:
                    processed_resume_data[key] = processed_resume_data.get(
                        key.lower().strip().replace(" ", "_").replace(":", ""), value
                    )

            # ── 5. Build flat lookup and render context ────────────────────
            flat_lookup = self._flatten_dict(processed_resume_data)
            render_context = {
                **processed_resume_data,
                "_": flat_lookup,
            }

            # ── 6. docxtpl render ──────────────────────────────────────────
            doc = DocxTemplate(processed_template_stream)
            doc.render(render_context)

            out_stream = io.BytesIO()
            doc.save(out_stream)
            
            return out_stream.getvalue(), json.dumps(linearized_data)

        except Exception as e:
            logger.error(f"Document rendering failed: {e}")
            raise RuntimeError(f"Failed to render document: {str(e)}")

    def prepare_document_markers(
        self,
        template_stream: io.BytesIO,
        field_list: List[str],
        manifest_list: Optional[List[Dict[str, Any]]] = None,
    ) -> io.BytesIO:
        """
        Scans the document for marker patterns (<<…>>, {{…}}, [[…]]) and rewrites
        them as docxtpl-compatible {{ _['fieldname'] }} tokens using manifest ordering.
        """
        doc = Document(template_stream)
        counter = 0
        manifest_list = manifest_list or []
        manifest_ptr = [0]

        MARKER_PATTERN = (
            r"<<\s*(.*?)\s*>>"
            r"|\{\{\s*(.*?)\s*\}\}"
            r"|\[\[\s*(.*?)\s*\]\]"
            r"|\[\s*(.*?)\s*\]"
            r"|<\s*(.*?)\s*>"
        )

        def _transform(text: str, fields: List[str], ctr: int, manifest: List, ptr: List) -> tuple:
            new_text = text
            offset = 0

            for match in list(re.finditer(MARKER_PATTERN, text)):
                start, end = match.span()
                original = match.group(0)
                raw = next((g for g in match.groups() if g is not None), "").strip()

                # ── manifest sequential match ──────────────────────────────
                if manifest and ptr[0] < len(manifest):
                    m = manifest[ptr[0]]
                    is_generic = "fill" in raw.lower() and "section" in raw.lower()
                    tag_match  = original == m.get("tag") or raw == m.get("tag")

                    if is_generic or tag_match:
                        target_key = m.get("fieldname") or m.get("meaning")
                        if target_key:
                            ptr[0] += 1
                            replacement = f"{{{{ _['{target_key}'] }}}}"
                            new_text = (
                                new_text[: start + offset]
                                + replacement
                                + new_text[end + offset :]
                            )
                            offset += len(replacement) - (end - start)
                            continue

                # ── legacy / exact field name match ───────────────────────
                if "fill" in raw.lower() and "section" in raw.lower():
                    replacement = f"{{{{ _['section_{ctr}'] }}}}"
                    new_text = new_text[: start + offset] + replacement + new_text[end + offset :]
                    offset += len(replacement) - len(original)
                    ctr += 1
                elif raw in fields:
                    replacement = f"{{{{ _['{raw}'] }}}}"
                    new_text = new_text[: start + offset] + replacement + new_text[end + offset :]
                    offset += len(replacement) - len(original)

            return new_text, ctr

        for p in doc.paragraphs:
            if "<<" in p.text or "{{" in p.text or "[[" in p.text:
                p.text, counter = _transform(p.text, field_list, counter, manifest_list, manifest_ptr)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "<<" in p.text or "{{" in p.text or "[[" in p.text:
                            p.text, counter = _transform(p.text, field_list, counter, manifest_list, manifest_ptr)

        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return out

    def generate_error_docx(self, template_id: str, error_message: str) -> bytes:
        """Produces a valid DOCX containing rendering failure details."""
        error_doc = Document()
        error_doc.add_heading("TEMPLATE RENDERING ERROR", level=1)
        error_doc.add_paragraph(f"Template Identification: {template_id}")
        error_doc.add_paragraph("-" * 20)
        error_doc.add_paragraph("Failure reason detected by system:")
        error_doc.add_paragraph(error_message)

        out = io.BytesIO()
        error_doc.save(out)
        return out.getvalue()

    def _flatten_dict(
        self,
        d: Dict[str, Any],
        parent_key: str = "",
        sep: str = ".",
    ) -> Dict[str, Any]:
        """Recursively flattens nested dicts to dot-notation keys."""
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._flatten_dict(v, new_key, sep=sep).items())
            else:
                items.append((new_key, v))
        return dict(items)
