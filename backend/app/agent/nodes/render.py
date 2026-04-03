from app.agent.state import AgentState
from app.services.resume_ai_service import ResumeAiService
import json
import os
import tempfile
import io
from docxtpl import DocxTemplate

def create_render_node(ai_service: ResumeAiService, storage):
    """
    Creates the LangGraph node for rendering the final outputs (resume summary and formatted CV).
    This node saves the artifacts to the StorageProvider and returns their URIs in the state.
    """
    async def render_node(state: AgentState) -> dict:
        print("Executing Render Node...")

        extracted_text = state.get("extracted_text", "")
        transformed_json_str = state.get("transformed_document_json", "")
        session_id = state.get("session_id", "unknown-session")

        # Use the centralized AI Service for summary generation
        summary_text = "Summary not available."
        resume_data = {}

        try:
            if transformed_json_str:
                resume_data = json.loads(transformed_json_str)

            if extracted_text:
                summary_guidance = state.get("summary_guidance") or ""
                industry = state.get("industry")
                language = state.get("language", "en")
                
                # Delegate to ResumeAiService for high-impact summary
                summary_text = await ai_service.generate_summary(
                    extracted_text=extracted_text,
                    guidance=summary_guidance,
                    industry=industry,
                    language=language
                )
            else:
                summary_text = "Original resume text not found. Summary cannot be generated."
        except Exception as e:
            print(f"Summary generation failed via AI Service: {e}")
            summary_text = "Summary generation failed or was empty."

        # Save summary with a consistent header
        summary_key = f"jobs/{session_id}/output/summary.md"
        final_summary_md = f"### CV Summary\n\n{summary_text}"
        summary_uri = storage.put_bytes(summary_key, final_summary_md.encode("utf-8"))

        template_id = state.get("selected_template_id") or "general_cv_v1"
        template_storage_uri = state.get("template_storage_uri")
        render_docx_uri = None

        # Fetch the template. Use template_storage_uri if resolved, else fallback to convention.
        try:
            if template_storage_uri:
                template_key = template_storage_uri.replace("local://", "")
            else:
                template_key = f"templates/{template_id}/template.docx"
            
            print(f"Retrieving template from storage using key: {template_key}")
            template_bytes = storage.get_bytes(template_key)
            template_stream = io.BytesIO(template_bytes)
            
            # Pre-process docx to convert << field >> to {{ field }} for docxtpl compatibility
            from docx import Document
            import re
            temp_doc = Document(template_stream)
            
            # Extract expected fields to handle sequential mapping of generic placeholders
            expected_fields_raw = state.get("expected_fields") or ""
            expected_fields_list = [f.strip() for f in expected_fields_raw.split(",") if f.strip()]
            generic_placeholder_count = 0
            
            # Simple paragraph and table cell replacement
            def replace_markers(text, field_list, counter):
                # Regex to find << ... >>
                matches = re.finditer(r'<<\s*(.*?)\s*>>', text)
                new_text = text
                offset = 0
                for match in matches:
                    original = match.group(0)
                    placeholder_content = match.group(1).strip().lower()
                    
                    # If it's a generic "fill this section" placeholder, try to map it sequentially
                    if "fill" in placeholder_content and "section" in placeholder_content and counter < len(field_list):
                        replacement = f"{{{{ {field_list[counter]} }}}}"
                        # print(f"Sequential mapping: '{original}' -> '{replacement}' (index {counter})")
                        counter += 1
                    else:
                        # Standard mapping: << summary >> -> {{ summary }}
                        replacement = f"{{{{ {placeholder_content} }}}}"
                    
                    start, end = match.span()
                    new_text = new_text[:start + offset] + replacement + new_text[end + offset:]
                    offset += len(replacement) - len(original)
                return new_text, counter
            
            # Process paragraphs
            for p in temp_doc.paragraphs:
                if '<<' in p.text:
                    p.text, generic_placeholder_count = replace_markers(p.text, expected_fields_list, generic_placeholder_count)
            
            # Process tables
            for table in temp_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if '<<' in p.text:
                                p.text, generic_placeholder_count = replace_markers(p.text, expected_fields_list, generic_placeholder_count)
            
            processed_stream = io.BytesIO()
            temp_doc.save(processed_stream)
            processed_stream.seek(0)
            
            doc = DocxTemplate(processed_stream)
            
            print(f"Rendering template {template_id} with data keys: {list(resume_data.keys())} and {generic_placeholder_count} mapped generic fields.")

            
            print(f"Rendering template {template_id} with data keys: {list(resume_data.keys())}")
            
            # Merge job metadata and summary into the rendering context
            render_context = {
                **resume_data,
                "summary": summary_text,
                "job_id": session_id
            }
            
            # Surface personal_info fields for easier template access (name, email)
            if "personal_info" in resume_data and isinstance(resume_data["personal_info"], dict):
                render_context.update(resume_data["personal_info"])
                
            doc.render(render_context)




            out_stream = io.BytesIO()
            doc.save(out_stream)
            out_stream.seek(0)

            render_key = f"jobs/{session_id}/output/formatted_resume.docx"
            render_docx_uri = storage.put_bytes(render_key, out_stream.read())
        except Exception as e:
            print(f"Failed to load or render template {template_id} from storage: {e}")
            render_key = f"jobs/{session_id}/output/formatted_resume.docx"
            fallback_text = f"TEMPLATE RENDERING ERROR\n\nTemplate: {template_id}\nError: {str(e)}"
            render_docx_uri = storage.put_bytes(render_key, fallback_text.encode("utf-8"))

        final_status = "rendered"
        if not state.get("validation_passed", True):
            final_status = "needs_review"

        return {
            "summary_text": summary_text,
            "summary_uri": summary_uri,
            "render_docx_uri": render_docx_uri,
            "status": final_status
        }


    return render_node
