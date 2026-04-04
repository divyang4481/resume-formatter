import os
import sys
import json
import asyncio
import logging
from typing import List, Dict, Any

# Ensure backend directory is in path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from docx import Document
from app.db.session import SessionLocal
from app.db.models import TemplateAsset
from app.dependencies import get_llm_runtime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("smart_ingester")

async def analyze_template_markers(template_path: str, template_id: str):
    logger.info(f"Starting High-IQ Smart Ingestion for Template: {template_id}")
    
    if not os.path.exists(template_path):
        logger.error(f"Template path not found: {template_path}")
        return

    llm = get_llm_runtime()
    doc = Document(template_path)
    
    # 1. Structure Capture: Get paragraphs with their context
    template_structure = []
    import re
    # BROAD marker detection to help the AI "Find" them
    MARKER_PATTERN = r"((?:<<|\{\{|\[\[)\s*(.*?)\s*(?:>>|\}\}|\]\]))"
    
    found_any = False
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if not p_text:
            continue
            
        markers = re.findall(MARKER_PATTERN, p_text)
        if markers:
            found_any = True
            template_structure.append({
                "type": "paragraph_with_markers",
                "text": p_text,
                "markers": [m[0] for m in markers]
            })
        else:
            # Context for headers/labels
            template_structure.append({
                "type": "context_paragraph",
                "text": p_text
            })

    # Also capture tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p_text = p.text.strip()
                    if not p_text: continue
                    markers = re.findall(MARKER_PATTERN, p_text)
                    if markers:
                        found_any = True
                        template_structure.append({
                            "type": "table_cell_with_markers",
                            "text": p_text,
                            "markers": [m[0] for m in markers]
                        })

    if not found_any:
        logger.warning(f"No markers found in {template_path}. Skipping AI analysis.")
        return

    # 2. HOLISTIC AI ANALYSIS: Provide the whole structure
    # We want the AI to understand the sequence and location
    prompt = f"""
    TASK: Analyze the structure of a Resume Word Template and map its placeholders to their SEMANTIC ROLES.
    
    FULL TEMPLATE STRUCTURE (SIMPLIFIED):
    {json.dumps(template_structure[:100], indent=2)}
    
    VALID SEMANTIC KEYS (Use these values in your mapping):
    - summary
    - professional_experience
    - key_projects
    - education
    - certifications
    - skills
    - personal_information.name
    - personal_information.email
    - personal_information.phone
    - personal_information.linkedin
    - personal_information.address
    
    INSTRUCTIONS:
    Identify EVERY marker you see (like "<<fill>>", "<<Section Fill>>", "{{address}}", etc.).
    Look at the HEADERS and text ABOVE each marker to decide its meaning.
    
    Example Logic: 
    - Marker "<<fill>>" is below "Professional Experience" header -> mapped to "professional_experience".
    - Marker "<<fill>>" is below "Education" header -> mapped to "education".
    
    OUTPUT: A JSON ARRAY of objects.
    Each object must have:
    - "tag": The EXACT raw marker string from the structure above.
    - "meaning": The VALID SEMANTIC KEY assigned.
    - "type": "section" (for large lists like jobs/edu) OR "field" (for single labels like name/address).
    
    OUTPUT ONLY VALID JSON.
    """
    
    logger.info("Executing High-IQ AI Analysis...")
    response = llm.generate(prompt=prompt, temperature=0.0)
    
    from app.agent.utils.llm_sanitizer import LlmSanitizer
    cleaned_json = LlmSanitizer.clean_json(response)
    
    try:
        manifest_list = json.loads(cleaned_json)
        logger.info(f"Generated High-IQ Manifest (length {len(manifest_list)})")
        logger.debug(f"Manifest contents: {json.dumps(manifest_list, indent=2)}")
        
        # 3. Persist to DB
        db = SessionLocal()
        try:
            asset = db.query(TemplateAsset).filter(TemplateAsset.id == template_id).first()
            if asset:
                # Save the rich JSON object array to the field_extraction_manifest column
                asset.field_extraction_manifest = json.dumps(manifest_list)
                
                # Sync expected_fields for backward compatibility
                all_meanings = list(set([m.get("meaning") for m in manifest_list if m.get("meaning")]))
                asset.expected_fields = ", ".join(all_meanings)
                
                db.commit()
                logger.info(f"SUCCESS: Template {template_id} now has a high-fidelity semantic manifest.")
            else:
                logger.error(f"Asset {template_id} not found in DB.")
        finally:
            db.close()
            
    except Exception as e:
        logger.error(f"High-IQ Analysis Failed to parse AI manifest: {e}")
        logger.error(f"Raw AI response: {response}")

if __name__ == "__main__":
    import asyncio
    import sys
    
    # Usage: python scripts/ingest_template_smart.py <template_id> <local_docx_path>
    if len(sys.argv) < 3:
        print("Usage: python scripts/ingest_template_smart.py <template_id> <local_docx_path>")
        sys.exit(1)
        
    tid = sys.argv[1]
    tpath = sys.argv[2]
    
    asyncio.run(analyze_template_markers(tpath, tid))
