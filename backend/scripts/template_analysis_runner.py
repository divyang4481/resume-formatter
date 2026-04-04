#!/usr/bin/env python3
"""Independent CV template analysis runner.

Usage:
  python template_analysis_runner.py /path/to/template.docx
  python template_analysis_runner.py /path/to/template.docx --output result.json
  python template_analysis_runner.py /path/to/template.docx --model llama3:latest --endpoint http://localhost:11434/api/generate

Requires:
  pip install httpx python-docx docxtpl

Notes:
  - This script is designed for DOCX templates.
  - It extracts visible text, reads undeclared Jinja/docxtpl variables,
    prompts an Ollama model with a strict JSON schema, validates the result,
    and prints the full JSON response.
"""

from __future__ import annotations

import argparse
import io
import json
import re
import sys
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from docxtpl import DocxTemplate


SYSTEM_PROMPT = """
You are a deterministic CV template structure analyzer.

Your job is to extract template guidance with strict structural fidelity.

Non-negotiable rules:
- Return JSON only
- Do not add markdown
- Do not add commentary
- Do not invent fields not visually implied by the template
- Do not merge multiple placeholders into one field
- Do not split one placeholder into multiple fields
- Preserve exact visual order of sections and placeholders
- The number of items in field_extraction_manifest must exactly equal the number of physical placeholders
- fieldname values must be high-level logical section names, not granular sub-fields
- expected_sections must contain only literal visible section headers in order
- If the template is ambiguous, choose the most standard professional interpretation without adding extra fields
- formatting_guidance must explain how downstream AI should format content inside each section without changing section structure
- summary_guidance must explain style, tone, density, and length for generated summaries

Before finalizing internally validate:
1. placeholder count equals manifest count
2. manifest order matches template order
3. output is valid JSON
4. no extra top-level keys are present
""".strip()


SCHEMA: dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "required": [
        "purpose",
        "expected_sections",
        "field_extraction_manifest",
        "summary_guidance",
        "formatting_guidance",
    ],
    "properties": {
        "purpose": {"type": "string"},
        "expected_sections": {"type": "string"},
        "field_extraction_manifest": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["fieldname", "meaning", "source_hints"],
                "properties": {
                    "fieldname": {"type": "string"},
                    "meaning": {"type": "string"},
                    "source_hints": {"type": "string"},
                },
            },
        },
        "summary_guidance": {"type": "string"},
        "formatting_guidance": {"type": "string"},
    },
}


EXAMPLE_BLOCK = {
    "purpose": "Professional resume template for technical candidates requiring a concise summary and a consolidated skills section.",
    "expected_sections": "Professional Summary, Technical Skills",
    "field_extraction_manifest": [
        {
            "fieldname": "professional_summary",
            "meaning": "Top summary section introducing candidate profile and value proposition.",
            "source_hints": "years of experience, role specialization, domain background, key strengths, major technologies",
        },
        {
            "fieldname": "technical_skills",
            "meaning": "Consolidated section for relevant technical competencies.",
            "source_hints": "programming languages, frameworks, cloud platforms, tools, databases",
        },
    ],
    "summary_guidance": "Write 3 to 5 lines in a recruiter-friendly tone. Emphasize seniority, domain fit, and strongest capabilities. Avoid personal pronouns and unsupported claims.",
    "formatting_guidance": "Use concise bullet grouping or compact categorized lists where appropriate. Keep content high-signal and avoid over-fragmenting into micro-sections unless the template explicitly requires it.",
}


def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract visible text from paragraphs and tables while preserving rough order."""
    doc = Document(io.BytesIO(docx_bytes))
    lines: list[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    cleaned: list[str] = []
    previous = None
    for line in lines:
        if line != previous:
            cleaned.append(line)
        previous = line

    return "\n".join(cleaned)


def extract_ordered_placeholders(docx_bytes: bytes) -> list[str]:
    """Extract docxtpl/Jinja placeholders with order preserved and duplicates removed."""
    tpl = DocxTemplate(io.BytesIO(docx_bytes))
    raw = [str(p).strip() for p in tpl.get_undeclared_template_variables()]
    ordered_unique = list(dict.fromkeys([p for p in raw if p]))
    return ordered_unique


def build_user_prompt(template_text: str, detected_placeholders: list[str]) -> str:
    placeholders_display = ", ".join(detected_placeholders) if detected_placeholders else "None identified programmatically."
    placeholder_count = len(detected_placeholders)

    return f"""
TASK: Reverse-engineer this CV template and produce governance metadata plus AI-steering guidance for downstream resume processing.

TEMPLATE TEXT:
{template_text}

DETECTED PLACEHOLDERS (IN ORDER):
{placeholders_display}

PLACEHOLDER COUNT:
{placeholder_count}

INSTRUCTIONS:
Analyze the template structure, visible headers, placeholder positions, and expected professional content.

Perform internally:
1. Extract visible headers in order
2. Count placeholders
3. Map each placeholder to exactly one high-level logical field
4. Preserve strict order
5. Write concise source_hints for each field
6. Write summary_guidance
7. Write formatting_guidance
8. Validate parity before output

STRICT RULES:
- expected_sections must be a comma-separated string of literal visible headers only
- field_extraction_manifest must contain exactly one item per placeholder
- the number of field_extraction_manifest items must equal PLACEHOLDER COUNT
- use the placeholder order exactly as provided
- do not create granular sub-fields when the template shows one larger section
- return only valid JSON
- output must conform exactly to the provided schema

Examples of acceptable fieldname values:
- professional_summary
- technical_skills
- domain_expertise
- professional_experience
- key_projects
- education
- certifications
- achievements

Examples of unacceptable fieldname values when the template only shows one larger section:
- programming_languages
- cloud_platforms
- tools_and_frameworks
- front_end_skills
- back_end_skills

EXAMPLE OUTPUT SHAPE:
{json.dumps(EXAMPLE_BLOCK, ensure_ascii=False)}

JSON SCHEMA TO FOLLOW:
{json.dumps(SCHEMA, ensure_ascii=False)}
""".strip()


def call_ollama(endpoint: str, model: str, system_prompt: str, user_prompt: str, temperature: float) -> str:
    payload = {
        "model": model,
        "system": system_prompt,
        "prompt": user_prompt,
        "format": SCHEMA,
        "stream": False,
        "raw": True,
        "options": {
            "temperature": temperature,
            "top_p": 0.9,
            "repeat_penalty": 1.05,
        },
    }

    with httpx.Client(timeout=300.0) as client:
        response = client.post(endpoint, json=payload)
        response.raise_for_status()
        data = response.json()
        return (data.get("response") or "").strip()


def try_parse_json(text: str) -> dict[str, Any]:
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}\s*$", text, flags=re.DOTALL)
        if match:
            return json.loads(match.group(0))
        raise


def validate_result(data: dict[str, Any], placeholder_count: int) -> None:
    required_top = {
        "purpose",
        "expected_sections",
        "field_extraction_manifest",
        "summary_guidance",
        "formatting_guidance",
    }
    extra = set(data.keys()) - required_top
    missing = required_top - set(data.keys())
    if missing:
        raise ValueError(f"Missing top-level keys: {sorted(missing)}")
    if extra:
        raise ValueError(f"Unexpected top-level keys: {sorted(extra)}")

    if not isinstance(data["field_extraction_manifest"], list):
        raise ValueError("field_extraction_manifest must be a list")

    manifest = data["field_extraction_manifest"]
    if len(manifest) != placeholder_count:
        raise ValueError(
            f"Placeholder parity failed: expected {placeholder_count} items, got {len(manifest)}"
        )

    for idx, item in enumerate(manifest, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"Manifest item {idx} must be an object")
        keys = set(item.keys())
        expected = {"fieldname", "meaning", "source_hints"}
        if keys != expected:
            raise ValueError(
                f"Manifest item {idx} keys invalid. Expected {sorted(expected)}, got {sorted(keys)}"
            )
        for key in expected:
            if not isinstance(item[key], str) or not item[key].strip():
                raise ValueError(f"Manifest item {idx}.{key} must be a non-empty string")

    for key in ["purpose", "expected_sections", "summary_guidance", "formatting_guidance"]:
        if not isinstance(data[key], str) or not data[key].strip():
            raise ValueError(f"{key} must be a non-empty string")


def run(template_path: Path, endpoint: str, model: str, temperature: float) -> dict[str, Any]:
    if not template_path.exists():
        raise FileNotFoundError(f"Template file not found: {template_path}")
    if template_path.suffix.lower() != ".docx":
        raise ValueError("This script currently supports only .docx templates")

    content = template_path.read_bytes()
    template_text = extract_docx_text(content)[:8000]
    detected_placeholders = extract_ordered_placeholders(content)

    user_prompt = build_user_prompt(template_text, detected_placeholders)
    raw_response = call_ollama(
        endpoint=endpoint,
        model=model,
        system_prompt=SYSTEM_PROMPT,
        user_prompt=user_prompt,
        temperature=temperature,
    )

    data = try_parse_json(raw_response)
    validate_result(data, placeholder_count=len(detected_placeholders))

    return {
        "template_path": str(template_path),
        "placeholder_count": len(detected_placeholders),
        "detected_placeholders": detected_placeholders,
        "template_text_preview": template_text,
        "analysis": data,
        "raw_llm_response": raw_response,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Analyze a DOCX CV template using local Ollama")
    parser.add_argument("template_path", help="Path to the DOCX template file")
    parser.add_argument("--model", default="llama3:latest", help="Ollama model name")
    parser.add_argument(
        "--endpoint",
        default="http://localhost:11434/api/generate",
        help="Ollama generate endpoint",
    )
    parser.add_argument("--temperature", type=float, default=0.1, help="Sampling temperature")
    parser.add_argument("--output", help="Optional path to save the full JSON result")
    args = parser.parse_args()

    try:
        result = run(
            template_path=Path(args.template_path),
            endpoint=args.endpoint,
            model=args.model,
            temperature=args.temperature,
        )
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    result_json = json.dumps(result, indent=2, ensure_ascii=False)
    if args.output:
        Path(args.output).write_text(result_json, encoding="utf-8")
        print(f"Saved result to {args.output}")
    else:
        print(result_json)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
