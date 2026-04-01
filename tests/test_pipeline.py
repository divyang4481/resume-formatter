"""Integration tests for the Pipeline."""

from __future__ import annotations

import io
import json
import textwrap

import pytest

from resume_formatter.pipeline import Pipeline, PipelineConfig, PipelineResult


SAMPLE_TEXT = textwrap.dedent("""\
    Eve Turner
    eve.turner@example.com
    +1 (800) 555-0199

    Summary
    Creative UI/UX designer with 5 years of experience.

    Experience
    Senior Designer at DesignCo
    June 2018 – Present
    Redesigned the flagship product, increasing user retention by 30%.

    Education
    Bachelor of Arts in Design
    Art University
    2018

    Skills
    Figma, Sketch, Adobe XD, CSS, HTML

    Languages
    English, German
""")


class TestPipelineRunText:
    """Tests that don't require actual document files."""

    def setup_method(self):
        self.pipeline = Pipeline()

    def test_returns_pipeline_result(self):
        result = self.pipeline.run_text(SAMPLE_TEXT)
        assert isinstance(result, PipelineResult)

    def test_resume_name_extracted(self):
        result = self.pipeline.run_text(SAMPLE_TEXT)
        assert result.resume.contact.name == "Eve Turner"

    def test_privacy_masks_email(self):
        result = self.pipeline.run_text(SAMPLE_TEXT)
        assert result.resume.contact.email != "eve.turner@example.com"

    def test_validation_report_populated(self):
        result = self.pipeline.run_text(SAMPLE_TEXT)
        assert result.validation is not None

    def test_privacy_report_populated(self):
        result = self.pipeline.run_text(SAMPLE_TEXT)
        assert result.privacy is not None
        assert result.privacy.has_pii

    def test_rendered_output_is_string(self):
        result = self.pipeline.run_text(SAMPLE_TEXT)
        assert isinstance(result.rendered, str)
        assert len(result.rendered) > 0

    def test_rendered_modern_html(self):
        config = PipelineConfig(template="modern")
        pipeline = Pipeline(config=config)
        result = pipeline.run_text(SAMPLE_TEXT)
        assert "<html" in result.rendered

    def test_rendered_json(self):
        config = PipelineConfig(template="json")
        pipeline = Pipeline(config=config)
        result = pipeline.run_text(SAMPLE_TEXT)
        data = json.loads(result.rendered)
        assert "contact" in data

    def test_rendered_minimal_markdown(self):
        config = PipelineConfig(template="minimal")
        pipeline = Pipeline(config=config)
        result = pipeline.run_text(SAMPLE_TEXT)
        assert "#" in result.rendered

    def test_privacy_disabled(self):
        config = PipelineConfig(apply_privacy=False)
        pipeline = Pipeline(config=config)
        result = pipeline.run_text(SAMPLE_TEXT)
        assert result.privacy is None
        assert result.resume.contact.email == "eve.turner@example.com"

    def test_validation_disabled(self):
        config = PipelineConfig(validate=False)
        pipeline = Pipeline(config=config)
        result = pipeline.run_text(SAMPLE_TEXT)
        assert result.validation is None

    def test_ok_property_true_for_valid_resume(self):
        result = self.pipeline.run_text(SAMPLE_TEXT)
        assert result.ok

    def test_raw_text_stored(self):
        result = self.pipeline.run_text(SAMPLE_TEXT)
        assert "Eve Turner" in result.raw_text


class TestPipelineRunFile:
    """Tests that use actual (in-memory) document files."""

    def test_run_docx(self, tmp_path):
        import io
        from docx import Document

        doc = Document()
        doc.add_paragraph("Frank Castle")
        doc.add_paragraph("frank@example.com")
        doc.add_paragraph("+1 555 001 0001")
        doc.add_paragraph("")
        doc.add_paragraph("Skills")
        doc.add_paragraph("Python, Go, Rust")
        buf = io.BytesIO()
        doc.save(buf)

        path = tmp_path / "frank.docx"
        path.write_bytes(buf.getvalue())

        result = Pipeline().run(path)
        assert isinstance(result, PipelineResult)
        assert result.resume.contact.name == "Frank Castle"

    def test_run_pdf(self, tmp_path):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        path = tmp_path / "blank.pdf"
        with open(path, "wb") as f:
            writer.write(f)

        # Blank PDF – should not crash; just return an empty (but valid) resume
        result = Pipeline().run(path)
        assert isinstance(result, PipelineResult)

    def test_unsupported_extension_raises(self, tmp_path):
        path = tmp_path / "file.xyz"
        path.write_text("data")
        with pytest.raises(ValueError, match="Unsupported file extension"):
            Pipeline().run(path)
